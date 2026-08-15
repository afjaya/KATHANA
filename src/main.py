import os
import json
import requests
from datetime import datetime
from docx import Document
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# Mengambil Environment Variables
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
GDRIVE_CREDENTIALS = os.environ.get("GDRIVE_CREDENTIALS")
GDRIVE_FOLDER_ID = os.environ.get("GDRIVE_FOLDER_ID")
DATA_PATH = "data/story.json"

def load_story_data():
    if not os.path.exists(DATA_PATH):
        raise FileNotFoundError(f"File {DATA_PATH} tidak ditemukan, Bosku!")
    with open(DATA_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def save_story_data(data):
    with open(DATA_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def generate_next_episode(story_data):
    bible = story_data["storyBible"]
    characters = story_data["characters"]
    episodes = story_data["episodes"]
    
    episode_number = len(episodes) + 1
    
    last_episode_summary = "Ini adalah awal mula cerita (Episode 1)."
    if episodes:
        last_ep = episodes[-1]
        last_episode_summary = f"Episode sebelumnya ({last_ep['title']}): {last_ep.get('summary', '')}"

    prompt = f"""
    Anda adalah penulis novel fantasi epik profesional. Buatlah **Episode {episode_number}** untuk semesta berikut:
    
    - Judul Universe: {bible['universeName']}
    - Genre: {bible['genre']}
    - Gaya Narasi: {bible['narratorStyle']}
    - Aturan Menulis: {bible['writingRules']}
    - Mekanisme Dunia: {json.dumps(bible['worldMechanics'], ensure_ascii=False)}
    - Karakter Aktif: {json.dumps(characters, ensure_ascii=False)}
    
    Kondisi Cerita Saat Ini:
    {last_episode_summary}
    
    Instruksi Khusus:
    - Panjang teks target sekitar {bible['episodeLength']} karakter.
    - Fokus pada ketegangan atmosferik, pertentangan antara Klan Lencana Perak Atas dan Bawah, serta penggunaan pedang perak dingin dan Dengung Giok.
    - Akhiri episode dengan ketegangan atau *cliffhanger* yang menarik untuk episode berikutnya.
    
    PENTING: Berikan output murni dalam format JSON object dengan tanda kurung kurawal pembuka {{ dan penutup }}. Jangan sertakan teks penjelasan lain di luar JSON.
    Struktur JSON:
    {{
      "title": "Judul Episode yang Menarik",
      "summary": "Ringkasan 1-2 kalimat dari peristiwa penting di episode ini",
      "content": "Teks isi cerita lengkap di sini..."
    }}
    """

    print(f"Sedang meracik Episode {episode_number} via REST API...")
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3.5-flash:generateContent?key={GEMINI_API_KEY}"
    
    payload = {
        "contents": [{
            "parts": [{"text": prompt}]
        }]
    }
    
    headers = {
        "Content-Type": "application/json"
    }
    
    response = requests.post(url, json=payload, headers=headers)
    
    if response.status_code != 200:
        raise Exception(f"Gagal memanggil API Google: {response.status_code} - {response.text}")
    
    res_data = response.json()
    
    try:
        text_response = res_data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError) as e:
        raise Exception(f"Struktur respons API tidak sesuai: {res_data}") from e
    
    # Bersihkan markdown block jika ada
    if "```json" in text_response:
        text_response = text_response.split("```json")[1].split("```")[0].strip()
    elif "```" in text_response:
        text_response = text_response.split("```")[1].split("```")[0].strip()
        
    new_ep_data = json.loads(text_response)
    
    new_episode = {
        "episodeNumber": episode_number,
        "title": new_ep_data.get("title", f"Episode {episode_number}"),
        "summary": new_ep_data.get("summary", ""),
        "content": new_ep_data.get("content", ""),
        "createdAt": datetime.utcnow().isoformat() + "Z"
    }
    
    return new_episode

def save_to_docx(episode):
    doc = Document()
    doc.add_heading(episode['title'], level=0)
    doc.add_paragraph(f"Episode: {episode['episodeNumber']} | Tanggal: {episode['createdAt']}")
    doc.add_heading("Ringkasan:", level=2)
    doc.add_paragraph(episode['summary'])
    doc.add_heading("Isi Cerita:", level=2)
    doc.add_paragraph(episode['content'])
    
    filename = f"stories/episode_{episode['episodeNumber']:03d}.docx"
    doc.save(filename)
    print(f"File dokumen {filename} berhasil dibuat secara lokal.")
    return filename

def upload_to_drive(file_path):
    if not GDRIVE_CREDENTIALS or not GDRIVE_FOLDER_ID:
        print("Peringatan: GDRIVE_CREDENTIALS atau GDRIVE_FOLDER_ID tidak diset. Lewat proses upload.")
        return
    
    print("Mengunggah dokumen ke Google Drive...")
    creds_json = json.loads(GDRIVE_CREDENTIALS)
    creds = service_account.Credentials.from_service_account_info(creds_json)
    service = build('drive', 'v3', credentials=creds)
    
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [GDRIVE_FOLDER_ID]
    }
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f"Berhasil! File terupload ke Google Drive dengan ID: {file.get('id')}")

if __name__ == "__main__":
    data = load_story_data()
    new_ep = generate_next_episode(data)
    
    data["episodes"].append(new_ep)
    data["storyBible"]["updatedAt"] = datetime.utcnow().isoformat() + "Z"
    
    save_story_data(data)
    
    # Buat file docx lalu upload ke Drive
    docx_file = save_to_docx(new_ep)
    upload_to_drive(docx_file)
    
    print(f"Selesai! Episode '{new_ep['title']}' sukses diproses sepenuhnya, Bosku.")