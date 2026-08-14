import os
import json
import docx
from dotenv import load_dotenv
from google import genai

# 1. Memuat Variabel dari file .env
load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY belum diset di file .env atau terminal!")

# Inisialisasi Client SDK google-genai terbaru
client = genai.Client(api_key=api_key)

# 2. Baca Database JSON
db_file = "db.json"
if not os.path.exists(db_file):
    raise FileNotFoundError(f"File '{db_file}' tidak ditemukan di direktori!")

with open(db_file, "r", encoding="utf-8") as f:
    db = json.load(f)

current_ep = db.get("currentEpisode", 0) + 1
story_bible = db.get("storyBible", {})
characters = db.get("characters", [])
last_summary = db.get("lastSummary", "")

# 3. Racik Prompt Dinamis (Membawa lastSummary dari episode sebelumnya)
prompt = f"""
You are a master fiction author writing a dark fantasy series.
Universe: {story_bible.get('universeName')}
Genre: {story_bible.get('genre')}
Style Rules: {story_bible.get('writingRules')}

Characters:
{json.dumps(characters, indent=2)}

Previous Episode Summary:
{last_summary if last_summary else "This is the very first episode. Introduce the world and characters as planned."}

Task: Write Episode #{current_ep} of the story.
Length: Approx {story_bible.get('episodeLength', 2000)} words.
Ensure the story is atmospheric, detailed, and ends on an engaging hook/cliffhanger.
"""

print(f"⚡ Sedang memproses Episode {current_ep} via Gemini AI...")

# 4. Panggil Gemini Model dengan Auto-Fallback (Mencegah Error 503)
candidate_models = [
    "gemini-3.5-flash",     # Pilihan Utama
    "gemini-2.5-pro",       # Cadangan 1 (Penulisan detail & dalam)
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-flash-latest"   # Cadangan 2 (Stabil)
]

response = None
used_model = None

for model_name in candidate_models:
    try:
        print(f"   Mengontak {model_name}...")
        response = client.models.generate_content(
            model=model_name,
            contents=prompt
        )
        used_model = model_name
        print(f"✅ Berhasil merespons menggunakan: {model_name}")
        break  # Keluar dari loop jika berhasil
    except Exception as e:
        print(f"⚠️ Model {model_name} kendala/busy. Mencoba model berikutnya...")

if not response:
    raise RuntimeError("Semua model cadangan gagal dihubungi. Silakan coba beberapa saat lagi.")

story_content = response.text

# 4b. Buat Ringkasan Episode (Summary Generation)
print(f"📝 Membuat ringkasan Episode {current_ep} untuk lastSummary...")
summary_prompt = f"Buatkan ringkasan singkat (2-3 kalimat) dalam bahasa Indonesia yang padat mengenai kejadian utama dan konflik yang terjadi pada episode cerita berikut:\n\n{story_content}"

try:
    summary_response = client.models.generate_content(
        model=used_model,
        contents=summary_prompt
    )
    episode_summary = summary_response.text.strip()
    print("✅ Ringkasan episode berhasil dibuat!")
except Exception as e:
    print(f"⚠️ Gagal membuat ringkasan otomatis: {e}")
    episode_summary = f"Episode {current_ep} telah selesai ditulis."

# 5. Simpan Hasil ke File .docx (Dipisah per Paragraf)
os.makedirs("stories", exist_ok=True)
doc = docx.Document()
doc.add_heading(f"Episode {current_ep}", level=1)

# Memisah teks per baris baru agar format di Word rapi
for paragraph in story_content.split("\n"):
    clean_p = paragraph.strip()
    if clean_p:
        doc.add_paragraph(clean_p)

docx_path = f"stories/Episode_{current_ep:03d}.docx"
doc.save(docx_path)
print(f"✅ Dokumen berhasil dibuat: {docx_path}")

# 6. Update Kembali db.json (Termasuk lastSummary)
db["currentEpisode"] = current_ep
db["lastSummary"] = episode_summary  # <--- Merekam ringkasan episode terbaru ke db.json

if "episodes" not in db:
    db["episodes"] = []

db["episodes"].append({
    "episodeNumber": current_ep,
    "title": f"Episode {current_ep}",
    "file": docx_path
})

with open(db_file, "w", encoding="utf-8") as f:
    json.dump(db, f, indent=2, ensure_ascii=False)

print("🎉 db.json berhasil diperbarui (termasuk lastSummary)! Mesin siap untuk episode berikutnya.")